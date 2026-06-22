"""Document export service."""
import os
from docx import Document

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
EXPORT_DIR = os.path.join(BASE_DIR, "data", "exports")
os.makedirs(EXPORT_DIR, exist_ok=True)


def export_sop(content: str, branding: dict = None) -> str:
    """Export SOP as a docx file. Returns file path."""
    doc = Document()

    lines = content.strip().split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    # Add branding footer
    if branding:
        doc.add_paragraph("")
        if branding.get("copyright"):
            doc.add_paragraph(branding["copyright"])
        if branding.get("signature"):
            doc.add_paragraph(f"作者: {branding['signature']}")

    filename = f"sop_{os.urandom(4).hex()}.docx"
    filepath = os.path.join(EXPORT_DIR, filename)
    doc.save(filepath)
    return filepath
