"""File parser: extract structured Markdown from uploaded files.

Supported formats:
  .docx  → python-docx (tables preserved)
  .xlsx  → openpyxl (each sheet as Markdown table)
  .csv   → csv module
  .pdf   → pdfplumber (tables preserved where detectable)
  .txt   → direct read
  .md    → direct read
  images → OCR placeholder (returns empty + status flag)
"""

import csv
import io
import os


def parse_file(filepath: str, filename: str = "") -> dict:
    """Parse a file and return {content_type, text, table_count, status}.

    Returns empty text with status='unsupported' for unrecognized formats.
    """
    ext = os.path.splitext(filename or filepath)[1].lower()

    if ext in (".txt", ".md", ".markdown"):
        return _parse_text(filepath)
    elif ext == ".csv":
        return _parse_csv(filepath)
    elif ext == ".docx":
        return _parse_docx(filepath)
    elif ext == ".xlsx":
        return _parse_xlsx(filepath)
    elif ext == ".pdf":
        return _parse_pdf(filepath)
    elif ext in (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"):
        return _parse_image(filepath)
    else:
        # Fallback: try as text
        try:
            return _parse_text(filepath)
        except Exception:
            return {"content_type": "markdown", "text": "", "table_count": 0,
                    "status": "unsupported", "error": f"Unsupported format: {ext}"}


def parse_bytes(data: bytes, filename: str) -> dict:
    """Parse file content from bytes (for API uploads)."""
    ext = os.path.splitext(filename)[1].lower()

    if ext in (".txt", ".md", ".markdown"):
        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError:
            text = data.decode("gbk", errors="ignore")
        return {"content_type": "markdown", "text": text, "table_count": 0, "status": "ok"}

    elif ext == ".csv":
        text = data.decode("utf-8", errors="ignore")
        return _csv_to_markdown(text)

    elif ext == ".docx":
        return _parse_docx_bytes(data)

    elif ext == ".xlsx":
        return _parse_xlsx_bytes(data)

    elif ext == ".pdf":
        return _parse_pdf_bytes(data)

    elif ext in (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"):
        return _parse_image_bytes(data, filename)

    else:
        try:
            text = data.decode("utf-8")
            return {"content_type": "markdown", "text": text, "table_count": 0, "status": "ok"}
        except Exception:
            return {"content_type": "markdown", "text": "", "table_count": 0,
                    "status": "unsupported", "error": f"Unsupported format: {ext}"}


# ── .txt / .md ──

def _parse_text(filepath: str) -> dict:
    for encoding in ("utf-8", "gbk", "latin-1"):
        try:
            with open(filepath, "r", encoding=encoding) as f:
                text = f.read()
            return {"content_type": "markdown", "text": text, "table_count": 0, "status": "ok"}
        except (UnicodeDecodeError, Exception):
            continue
    return {"content_type": "markdown", "text": "", "table_count": 0,
            "status": "error", "error": "Failed to decode text file"}


# ── .csv ──

def _parse_csv(filepath: str) -> dict:
    for encoding in ("utf-8", "gbk", "latin-1"):
        try:
            with open(filepath, "r", encoding=encoding, newline="") as f:
                text = f.read()
            return _csv_to_markdown(text)
        except (UnicodeDecodeError, Exception):
            continue
    return {"content_type": "markdown", "text": "", "table_count": 0,
            "status": "error", "error": "Failed to decode CSV file"}


def _csv_to_markdown(text: str) -> dict:
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)
    if not rows:
        return {"content_type": "markdown", "text": "", "table_count": 0, "status": "ok"}

    md = _rows_to_md_table(rows)
    return {"content_type": "markdown", "text": md, "table_count": 1, "status": "ok"}


# ── .docx ──

def _parse_docx(filepath: str) -> dict:
    try:
        from docx import Document
    except ImportError:
        return {"content_type": "markdown", "text": "", "table_count": 0,
                "status": "error", "error": "python-docx not installed"}

    try:
        doc = Document(filepath)
        return _docx_to_markdown(doc)
    except Exception as e:
        return {"content_type": "markdown", "text": "", "table_count": 0,
                "status": "error", "error": str(e)}


def _parse_docx_bytes(data: bytes) -> dict:
    try:
        from docx import Document
    except ImportError:
        return {"content_type": "markdown", "text": "", "table_count": 0,
                "status": "error", "error": "python-docx not installed"}

    try:
        doc = Document(io.BytesIO(data))
        return _docx_to_markdown(doc)
    except Exception as e:
        return {"content_type": "markdown", "text": "", "table_count": 0,
                "status": "error", "error": str(e)}


def _docx_to_markdown(doc) -> dict:
    parts = []
    table_count = 0

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
        if tag == "p":
            para = _find_paragraph(doc, element)
            if para:
                text = para.text.strip()
                if text:
                    style = para.style.name if para.style else ""
                    if "Heading" in style or "heading" in style or "标题" in style:
                        level = 1
                        for c in style:
                            if c.isdigit():
                                level = int(c)
                                break
                        parts.append(f"{'#' * level} {text}\n")
                    else:
                        parts.append(f"{text}\n")
        elif tag == "tbl":
            table = _find_table(doc, element)
            if table:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(cells)
                if rows:
                    parts.append(_rows_to_md_table(rows))
                    parts.append("")
                    table_count += 1

    text = "\n".join(parts)
    return {"content_type": "markdown", "text": text, "table_count": table_count, "status": "ok"}


def _find_paragraph(doc, element):
    for para in doc.paragraphs:
        if para._element is element:
            return para
    return None


def _find_table(doc, element):
    for table in doc.tables:
        if table._element is element:
            return table
    return None


# ── .xlsx ──

def _parse_xlsx(filepath: str) -> dict:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return {"content_type": "markdown", "text": "", "table_count": 0,
                "status": "error", "error": "openpyxl not installed"}

    try:
        wb = load_workbook(filepath, data_only=True)
        return _xlsx_to_markdown(wb)
    except Exception as e:
        return {"content_type": "markdown", "text": "", "table_count": 0,
                "status": "error", "error": str(e)}


def _parse_xlsx_bytes(data: bytes) -> dict:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return {"content_type": "markdown", "text": "", "table_count": 0,
                "status": "error", "error": "openpyxl not installed"}

    try:
        wb = load_workbook(io.BytesIO(data), data_only=True)
        return _xlsx_to_markdown(wb)
    except Exception as e:
        return {"content_type": "markdown", "text": "", "table_count": 0,
                "status": "error", "error": str(e)}


def _xlsx_to_markdown(wb) -> dict:
    parts = []
    table_count = 0

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            rows.append([str(c) if c is not None else "" for c in row])
        if not rows:
            continue

        if len(wb.sheetnames) > 1:
            parts.append(f"## {sheet_name}\n")
        parts.append(_rows_to_md_table(rows))
        parts.append("")
        table_count += 1

    text = "\n".join(parts)
    return {"content_type": "markdown", "text": text, "table_count": table_count, "status": "ok"}


# ── .pdf ──

def _parse_pdf(filepath: str) -> dict:
    try:
        import pdfplumber
    except ImportError:
        return {"content_type": "markdown", "text": "", "table_count": 0,
                "status": "error", "error": "pdfplumber not installed"}

    try:
        parts = []
        table_count = 0
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                # Extract tables first
                tables = page.extract_tables()
                for tbl in tables:
                    if tbl:
                        parts.append(_rows_to_md_table(tbl))
                        parts.append("")
                        table_count += 1

                # Extract text (non-table content)
                text = page.extract_text()
                if text:
                    parts.append(text)
                    parts.append("")

        text = "\n".join(parts)
        return {"content_type": "markdown", "text": text, "table_count": table_count, "status": "ok"}
    except Exception as e:
        return {"content_type": "markdown", "text": "", "table_count": 0,
                "status": "error", "error": str(e)}


def _parse_pdf_bytes(data: bytes) -> dict:
    try:
        import pdfplumber
    except ImportError:
        return {"content_type": "markdown", "text": "", "table_count": 0,
                "status": "error", "error": "pdfplumber not installed"}

    try:
        parts = []
        table_count = 0
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                for tbl in tables:
                    if tbl:
                        parts.append(_rows_to_md_table(tbl))
                        parts.append("")
                        table_count += 1

                text = page.extract_text()
                if text:
                    parts.append(text)
                    parts.append("")

        text = "\n".join(parts)
        return {"content_type": "markdown", "text": text, "table_count": table_count, "status": "ok"}
    except Exception as e:
        return {"content_type": "markdown", "text": "", "table_count": 0,
                "status": "error", "error": str(e)}


# ── Images (OCR placeholder) ──

def _parse_image(filepath: str) -> dict:
    return {"content_type": "markdown", "text": "", "table_count": 0,
            "status": "unsupported",
            "error": "OCR not yet implemented. Image files require OCR API integration."}


def _parse_image_bytes(data: bytes, filename: str) -> dict:
    return {"content_type": "markdown", "text": "", "table_count": 0,
            "status": "unsupported",
            "error": "OCR not yet implemented. Image files require OCR API integration."}


# ── Helpers ──

def _rows_to_md_table(rows: list) -> str:
    """Convert list of lists to a Markdown table."""
    if not rows:
        return ""

    # Normalize all cells to strings and trim
    clean = [[str(c).strip().replace("\n", "<br>") for c in row] for row in rows]

    # Pad rows to the width of the widest row
    max_cols = max(len(row) for row in clean) if clean else 0
    padded = [row + [""] * (max_cols - len(row)) for row in clean]

    lines = []
    # Header row
    lines.append("| " + " | ".join(padded[0]) + " |")
    # Separator
    lines.append("| " + " | ".join(["---"] * max_cols) + " |")
    # Data rows
    for row in padded[1:]:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)
